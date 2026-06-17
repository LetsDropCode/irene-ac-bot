from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUT = DOCS / "Irene_AC_Bot_Technical_Workbook_and_User_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(100, 100, 100)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DFF4E8"
PHONE_BG = (239, 242, 245)
WHATSAPP_GREEN = (37, 211, 102)


def font(size=28, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def wrapped_lines(draw, text, max_width, fnt):
    lines = []
    for para in text.split("\n"):
        if not para:
            lines.append("")
            continue
        words = para.split(" ")
        line = ""
        for word in words:
            test = word if not line else f"{line} {word}"
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = word
        if line:
            lines.append(line)
    return lines


def round_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_phone_mock(filename, title, messages, footer=None):
    width, height = 900, 1300
    image = Image.new("RGB", (width, height), PHONE_BG)
    draw = ImageDraw.Draw(image)
    round_rect(draw, (55, 35, width - 55, height - 35), 48, (255, 255, 255), (210, 216, 222), 3)
    round_rect(draw, (55, 35, width - 55, 155), 48, WHATSAPP_GREEN)
    draw.rectangle((55, 105, width - 55, 155), fill=WHATSAPP_GREEN)
    draw.text((105, 76), title, fill=(255, 255, 255), font=font(34, True))
    draw.text((105, 118), "Mock WhatsApp conversation", fill=(236, 255, 242), font=font(18))

    y = 200
    body_font = font(28)
    small_font = font(22)
    for side, text, buttons in messages:
        bubble_width = 610
        x = 105 if side == "bot" else width - 105 - bubble_width
        fill = (255, 255, 255) if side == "bot" else (220, 248, 198)
        lines = wrapped_lines(draw, text, bubble_width - 56, body_font)
        line_h = 38
        bubble_h = 36 + len(lines) * line_h + (56 if buttons else 0)
        round_rect(draw, (x, y, x + bubble_width, y + bubble_h), 24, fill, (220, 225, 230), 1)
        ty = y + 22
        for line in lines:
            draw.text((x + 28, ty), line, fill=(35, 35, 35), font=body_font)
            ty += line_h
        if buttons:
            bx = x + 28
            by = y + bubble_h - 46
            for label in buttons:
                bw = min(175, max(110, draw.textbbox((0, 0), label, font=small_font)[2] + 36))
                round_rect(draw, (bx, by, bx + bw, by + 34), 14, (245, 248, 250), (184, 202, 214), 1)
                draw.text((bx + 18, by + 6), label, fill=(23, 105, 170), font=small_font)
                bx += bw + 14
        y += bubble_h + 28

    if footer:
        draw.text((105, height - 100), footer, fill=(90, 90, 90), font=font(21))

    path = ASSETS / filename
    image.save(path)
    return path


def draw_admin_mock():
    width, height = 1100, 720
    image = Image.new("RGB", (width, height), (247, 249, 252))
    draw = ImageDraw.Draw(image)
    draw.text((60, 48), "Irene AC Bot - Admin status mockup", fill=(20, 42, 70), font=font(38, True))
    draw.text((62, 96), "Example operator view derived from TT STATUS, PENDING and scheduled job outputs.", fill=(90, 99, 110), font=font(22))

    cards = [
        ("TT window", "Tuesday", "Closes 22:30"),
        ("Checked in", "42 members", "+8 pending"),
        ("Runner results", "31 confirmed", "4/6/8 km"),
        ("Walker activity", "7 logged", "Opt-out honored"),
    ]
    x = 60
    for title, value, note in cards:
        round_rect(draw, (x, 150, x + 235, 275), 18, (255, 255, 255), (218, 225, 232), 2)
        draw.text((x + 20, 172), title, fill=(31, 77, 120), font=font(22, True))
        draw.text((x + 20, 208), value, fill=(25, 25, 25), font=font(25, True))
        draw.text((x + 20, 244), note, fill=(90, 99, 110), font=font(20))
        x += 255

    round_rect(draw, (60, 320, 1040, 650), 16, (255, 255, 255), (218, 225, 232), 2)
    headers = ["Command", "Purpose", "Expected response"]
    col_x = [90, 330, 690]
    for h, cx in zip(headers, col_x):
        draw.text((cx, 350), h, fill=(31, 77, 120), font=font(23, True))
    rows = [
        ("TT CODE", "Get tonight's code", "Admins only"),
        ("TT STATUS", "Check participation health", "Counts and pending status"),
        ("PENDING", "Find missing submissions", "Names and phone numbers"),
        ("RECOVER TONIGHT", "Resend prompts", "Counts by runner/walker/both"),
        ("SEASON", "Season PB leaderboard", "Ranked by distance"),
    ]
    y = 395
    for row in rows:
        draw.line((85, y - 10, 1015, y - 10), fill=(232, 236, 240), width=1)
        for val, cx in zip(row, col_x):
            draw.text((cx, y), val, fill=(40, 45, 50), font=font(21))
        y += 48
    path = ASSETS / "mock-admin-status.png"
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(text)
    doc.add_paragraph()


def add_image(doc, path, caption, width=5.65):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = MUTED


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, DARK_BLUE, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def write_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    onboarding = draw_phone_mock(
        "mock-onboarding.png",
        "Irene AC Bot",
        [
            ("user", "Hi", []),
            ("bot", "Reply OK to continue or STOP to opt out.", []),
            ("user", "OK", []),
            ("bot", "Send your first and last name.", []),
            ("user", "Naledi Mokoena", []),
            ("bot", "How do you usually participate?", ["Runner", "Walker", "Both"]),
        ],
        "Example only - member names and phone numbers are fictional.",
    )
    submission = draw_phone_mock(
        "mock-submission.png",
        "Submit TT result",
        [
            ("user", "Submit TT", []),
            ("bot", "Send tonight's TT code to check in.", []),
            ("user", "TT-4821", []),
            ("bot", "Checked in! Select your TT distance:", ["4 km", "6 km", "8 km"]),
            ("user", "6 km", []),
            ("bot", "Send your time.", []),
            ("user", "27:41", []),
            ("bot", "Please confirm your Time Trial:\nDistance: 6 km\nTime: 27:41", ["Confirm", "Edit"]),
        ],
    )
    profile = draw_phone_mock(
        "mock-profile-progress.png",
        "Profile and progress",
        [
            ("user", "Progress", []),
            ("bot", "Naledi, your progress\nTT activities: 8\nLatest: 6km - 27:41 (4:37/km)\nNext milestone: 10 activities (2 to go)\nPBs\n6km - 27:41", []),
            ("user", "Profile", []),
            ("bot", "Your TT Profile\nName: Naledi Mokoena\nParticipation: Runner\nRuns: 8\nPersonal Bests\n6km - 27:41", ["Edit name", "Change type"]),
        ],
    )
    leaderboard = draw_phone_mock(
        "mock-leaderboard.png",
        "Leaderboard",
        [
            ("user", "Leaderboard", []),
            ("bot", "Tonight's TT Leaderboard\n\n8 km\n1. A. Runner - 36:12\n2. B. Runner - 37:05\n\n6 km\n1. Naledi Mokoena - 27:41\n2. Sam Patel - 28:10\n\nWalker Activity\nThandi Dlamini - brisk 45 min walk", []),
        ],
    )
    admin = draw_admin_mock()

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Irene AC Bot").bold = True
    subtitle = doc.add_paragraph("Technical Workbook and User Guide")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = DARK_BLUE
    meta = doc.add_paragraph("Prepared for operators, administrators, and maintainers | Generated from the local codebase")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = MUTED
    doc.add_paragraph()
    add_note(
        doc,
        "Scope",
        "This guide documents the FastAPI WhatsApp bot in this repository. Mock screenshots are illustrative examples, not live user data.",
    )
    doc.add_page_break()

    doc.add_heading("1. Quick Start", level=1)
    doc.add_paragraph(
        "Irene AC Bot helps Irene Athletics Club members check in for Tuesday time trials, submit results or walking activity, view profiles and progress, and receive leaderboards. Admin users can generate TT codes, inspect status, list pending submissions, recover prompts, and view season PB leaderboards."
    )
    add_table(
        doc,
        ["Area", "What to know"],
        [
            ("Main app", "FastAPI app in app/main.py with root, health, and WhatsApp webhook routes."),
            ("Messaging", "WhatsApp Cloud API sender helpers live in app/whatsapp.py."),
            ("Database", "PostgreSQL via psycopg2. Startup runs idempotent table setup and safe migrations."),
            ("TT window", "Submissions are allowed on Tuesdays and close at 22:30 Africa/Johannesburg time."),
            ("Scheduled job", "Wednesday 08:00 leaderboard broadcast, documented in docs/scheduled-jobs.md."),
        ],
        [1.6, 4.9],
    )

    doc.add_heading("2. User Guide", level=1)
    doc.add_heading("2.1 First-time onboarding", level=2)
    add_numbers(
        doc,
        [
            "The member starts by sending HELP, MENU, START, HI, HELLO, /HELP, or any supported greeting.",
            "The bot asks for POPIA acknowledgement. The member replies OK to continue or STOP to opt out.",
            "The bot asks for first and last name.",
            "The member chooses Runner, Walker, or Both.",
            "The bot stores the profile and the member can submit TT results when the TT window is open.",
        ],
    )
    add_image(doc, onboarding, "Mock screenshot 1: first-time onboarding flow.", 4.9)

    doc.add_heading("2.2 Submitting a runner TT result", level=2)
    add_numbers(
        doc,
        [
            "Open the menu or type Submit TT, TT, or Submit.",
            "Send the current TT code supplied by an admin.",
            "Choose 4 km, 6 km, or 8 km.",
            "Send the time in mm:ss or hh:mm:ss format.",
            "Review the confirmation card and tap Confirm, or tap Edit to re-enter the result.",
        ],
    )
    add_note(doc, "Validation", "Times must use valid minutes and seconds. Examples: 27:41 or 01:27:41.")
    add_image(doc, submission, "Mock screenshot 2: runner result submission and confirmation.", 4.9)

    doc.add_heading("2.3 Walkers and Both participants", level=2)
    add_bullets(
        doc,
        [
            "Walkers are prompted to describe their workout instead of choosing a race distance.",
            "Both participants choose whether they are submitting a distance result or a walking workout.",
            "Walking submissions are confirmed immediately after the workout description is saved.",
        ],
    )

    doc.add_heading("2.4 Profile, progress, and leaderboard", level=2)
    add_bullets(
        doc,
        [
            "Profile shows saved name, participation type, total runs, personal bests, and recent results.",
            "Progress shows latest activity, next milestone, PBs, and trend messaging where enough history exists.",
            "Leaderboard shows current TT runner rankings by distance plus recent walking activity.",
            "Members can opt out of leaderboard sharing by typing STOP, OPT OUT, STOP LEADERBOARD, or choosing Stop sharing.",
        ],
    )
    add_image(doc, profile, "Mock screenshot 3: profile and progress examples.", 4.9)
    add_image(doc, leaderboard, "Mock screenshot 4: member leaderboard view.", 4.9)

    doc.add_heading("3. Admin Guide", level=1)
    add_table(
        doc,
        ["Command", "Purpose", "Notes"],
        [
            ("TT CODE", "Generate or retrieve tonight's code.", "Admin only. Also available from the admin menu."),
            ("TT STATUS", "Show current TT participation and pending state.", "Useful during TT evenings."),
            ("PENDING", "List checked-in members still missing completed submissions.", "Includes names and phone numbers."),
            ("RECOVER TONIGHT", "Resend submission prompts to checked-in members with no result details.", "Counts runners, walkers, and both participants."),
            ("SEASON", "Show season personal-best leaderboard.", "Ranks by distance and best time."),
            ("LEADERBOARD", "Show tonight's full leaderboard.", "Available to admins and members."),
        ],
        [1.35, 2.75, 2.4],
    )
    add_image(doc, admin, "Mock screenshot 5: operator status dashboard concept based on bot outputs.", 6.1)

    doc.add_heading("4. Technical Workbook", level=1)
    doc.add_heading("4.1 Application architecture", level=2)
    add_table(
        doc,
        ["Layer", "Files", "Responsibility"],
        [
            ("HTTP app", "app/main.py, app/webhook.py", "FastAPI startup, health endpoints, webhook verification and message routing."),
            ("WhatsApp API", "app/whatsapp.py", "Builds text, list, and button payloads for WhatsApp Cloud API v22.0."),
            ("Database", "app/db.py", "PostgreSQL connection helpers, startup schema creation, migrations, indexes, event seed data."),
            ("Flows", "app/flows/*", "Help/menu command resolution and pending submission state handling."),
            ("Services", "app/services/*", "Members, submissions, attendance, profile/progress, leaderboards, validation, TT status, OpenAI coaching."),
            ("Jobs", "scripts/*", "Scheduled leaderboard and profile completion campaigns."),
            ("Tests", "tests/*", "Unit and integration-style tests for webhook flow, services, formatting, and OpenAI fallback behavior."),
        ],
        [1.35, 2.25, 2.9],
    )

    doc.add_heading("4.2 Environment variables", level=2)
    add_table(
        doc,
        ["Variable", "Required", "Purpose"],
        [
            ("DATABASE_URL", "Yes", "PostgreSQL connection string. The app fails fast if missing."),
            ("VERIFY_TOKEN", "Yes for Meta setup", "Webhook verification token used by WhatsApp/Meta."),
            ("WHATSAPP_TOKEN", "Yes for sends", "Bearer token for WhatsApp Cloud API."),
            ("PHONE_NUMBER_ID", "Yes for sends", "Meta phone number ID used in Graph API message URL."),
            ("ENV", "No", "Returned by the root endpoint; defaults to development."),
            ("WHATSAPP_CONNECT_TIMEOUT", "No", "HTTP connect timeout. Defaults to 2 seconds."),
            ("WHATSAPP_READ_TIMEOUT", "No", "HTTP read timeout. Defaults to 5 seconds."),
            ("OPENAI_API_KEY", "Needed for coaching", "Used by the OpenAI coaching service where configured."),
        ],
        [1.9, 1.05, 3.55],
    )

    doc.add_heading("4.3 Local setup checklist", level=2)
    add_numbers(
        doc,
        [
            "Create and activate a Python virtual environment.",
            "Install dependencies from requirements.txt.",
            "Create a .env file with DATABASE_URL, VERIFY_TOKEN, WHATSAPP_TOKEN, PHONE_NUMBER_ID, and any OpenAI configuration.",
            "Start the app with uvicorn app.main:app --reload for local development.",
            "Confirm /health returns status ok.",
            "Configure the public webhook URL in Meta when deploying.",
        ],
    )
    add_note(doc, "Deployment", "The Procfile starts uvicorn on host 0.0.0.0 and the platform-provided PORT, which matches Railway-style deployment.")

    doc.add_heading("4.4 Data model", level=2)
    add_table(
        doc,
        ["Table", "Key fields", "Role"],
        [
            ("members", "phone, first_name, last_name, participation_type, profile_state, popia_acknowledged, leaderboard_opt_out", "Stores member identity, preferences, consent state, and temporary profile edit state."),
            ("submissions", "member_id, activity, distance_text, time_text, seconds, status, tt_code, tt_code_verified, confirmed", "Stores daily TT or walking activity submissions and their completion state."),
            ("event_codes", "event, code, event_date", "Stores valid TT codes for the current event date."),
            ("event_config", "event, day_of_week, open_time, close_time, active", "Seeds event metadata for TT, WEDLSD, and SUNSOCIAL."),
            ("attendance", "member_id, event, event_date, source", "Tracks checked-in members and prevents duplicate attendance rows."),
        ],
        [1.25, 2.9, 2.35],
    )

    doc.add_heading("4.5 Webhook state machine", level=2)
    add_numbers(
        doc,
        [
            "Extract sender, text body, or interactive button/list reply from the WhatsApp payload.",
            "Handle HELP and menu actions before normal submission handling.",
            "Run admin commands only when the sender phone number is in the admin allow-list.",
            "Create a member if the sender is new.",
            "Require POPIA acknowledgement and complete profile basics before TT submission.",
            "Create or reuse today's non-cancelled submission.",
            "Gate TT activity to the allowed Tuesday window.",
            "Verify TT code, mark attendance, collect distance or workout, collect time, then confirm.",
            "After confirmation, send a recap with pace, PB/milestone information, rank, and optional coaching feedback.",
        ],
    )

    doc.add_heading("4.6 Scheduled operations", level=2)
    add_table(
        doc,
        ["Job", "Schedule", "What it does"],
        [
            ("Next-day TT leaderboard", "Wednesday 08:00 Africa/Johannesburg", "Runs scripts/send_next_day_leaderboard.py and sends yesterday's TT leaderboard to checked-in members who have not opted out."),
            ("Profile completion campaign", "As needed", "scripts/profile_completion_campaign.py can prompt members with incomplete placeholder profiles."),
        ],
        [2.1, 1.8, 2.6],
    )

    doc.add_heading("4.7 Test and release checklist", level=2)
    add_bullets(
        doc,
        [
            "Run the existing test suite before deployment.",
            "Check webhook state-flow tests when changing message routing.",
            "Check WhatsApp menu tests when changing button/list IDs or menu wording.",
            "Check leaderboard and progress formatter tests when changing result display.",
            "Confirm environment variables are present before deploying.",
            "Perform a live admin smoke test: TT CODE, TT STATUS, PENDING, and /health.",
        ],
    )

    doc.add_heading("5. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Resolution"],
        [
            ("App fails at startup", "DATABASE_URL is missing or unreachable.", "Set DATABASE_URL and confirm the database accepts connections."),
            ("Messages are received but no replies send", "WHATSAPP_TOKEN or PHONE_NUMBER_ID missing/expired.", "Refresh Meta token and confirm phone number ID."),
            ("Members cannot submit", "Not Tuesday, after 22:30, invalid TT code, or code missing for today's date.", "Check TT window and event_codes table."),
            ("Time rejected", "Format does not match mm:ss or hh:mm:ss, or minutes/seconds are 60 or higher.", "Ask member to resend an example like 27:41."),
            ("Admin command ignored", "Sender number is not in the admin allow-list.", "Add the phone number in app/webhook.py and app/config.py, then redeploy."),
            ("Leaderboard missing someone", "Submission not COMPLETE, distance missing for runners, or member opted out.", "Inspect submissions status and leaderboard_opt_out."),
        ],
        [1.85, 2.2, 2.45],
    )

    doc.add_heading("6. Workbook Exercises", level=1)
    add_table(
        doc,
        ["Exercise", "Goal", "Done when"],
        [
            ("Onboard a test member", "Verify POPIA, name capture, and participation type.", "Profile command shows the correct name and type."),
            ("Submit a 6 km result", "Verify TT code, distance, time validation, confirmation, and recap.", "Submission is COMPLETE and leaderboard includes the result."),
            ("Submit a walker activity", "Verify non-distance workflow.", "Walker activity appears in the walking feed."),
            ("Run admin recovery", "Confirm prompt resend path.", "RECOVER TONIGHT reports counts by participation type."),
            ("Simulate opt-out", "Verify privacy preference.", "Member no longer receives scheduled leaderboard broadcasts."),
            ("Change profile details", "Verify edit_name and edit_type buttons.", "Updated profile persists after another message."),
        ],
        [1.7, 2.6, 2.2],
    )

    doc.add_heading("7. Maintenance Notes", level=1)
    add_bullets(
        doc,
        [
            "Keep admin phone numbers synchronized wherever admin allow-lists are defined.",
            "Treat WhatsApp interactive IDs as contracts: changing them requires corresponding changes in help_flow.py and tests.",
            "Prefer adding database changes as idempotent ALTER TABLE or CREATE INDEX statements inside init_db.",
            "Keep all date logic explicit about Africa/Johannesburg time.",
            "Avoid logging secrets or full access tokens when debugging WhatsApp send failures.",
            "Refresh mock screenshots and this workbook after major flow changes.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Irene AC Bot technical workbook and user guide")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(write_doc())
